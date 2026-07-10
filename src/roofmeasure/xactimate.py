"""Modular Xactimate-style estimate generator (.docx).

Python port of the Stronghouse Solutions docx-js engine. Fully config-driven:
company name/logo, estimator/inspector, homeowner, claim info, tax rate,
price list, O&P toggle, any number of structures, arbitrary line items.

Line item shape: (description, qty, unit, unit_price, material_fraction, note?)
Material sales tax applies only to the material fraction of each line.
"""

from __future__ import annotations

import os
from dataclasses import dataclass, field

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ------------------------------------------------------------------ config

@dataclass
class EstimateConfig:
    out_file: str
    # letterhead
    company_name: str = "Stronghouse Solutions"
    company_tagline: str = "Roofing & Exteriors"
    company_city: str = ""
    logo_path: str = ""            # optional PNG
    # parties
    owner: str = "TBD"
    phone: str = "TBD"
    property_lines: list[str] = field(default_factory=list)
    estimator_name: str = "Stronghouse Solutions"
    contractor_company: str = "Stronghouse Solutions"
    contractor_city: str = ""
    # claim
    claim_number: str = "TBD"
    policy_number: str = "TBD"
    type_of_loss: str = "Wind/Hail"
    date_of_loss: str = "TBD"
    date_inspected: str = ""
    date_entered: str = ""
    # estimate meta
    price_list: str = ""
    estimate_name: str = ""
    footer_date: str = ""
    tax_rate: float = 0.0825
    include_op: bool = False       # Stronghouse default: no O&P unless requested
    op_overhead_pct: float = 0.10
    op_profit_pct: float = 0.10
    structures: list[dict] = field(default_factory=list)
    # structure dict: {heading, summary_heading, totals_label, recap_label,
    #                  measurements: [(value, label), ...],
    #                  items: [(desc, qty, unit, price, matl_frac, note?), ...]}


def money(n: float) -> str:
    neg = n < 0
    v = f"{abs(n):,.2f}"
    return f"({v})" if neg else v


# ------------------------------------------------------------------ docx helpers

FONT = "Times New Roman"


def _run(p, text: str, size: float = 10, bold: bool = False, italic: bool = False):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return r


def _cell_borders(cell, top: bool = False, bottom: bool = False):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for tag, on in (("w:top", top), ("w:bottom", bottom)):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single" if on else "nil")
        if on:
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), "000000")
        borders.append(el)
    tcPr.append(borders)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _tight(p):
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    return p


def _para_rule(doc, double: bool = False):
    p = doc.add_paragraph()
    _tight(p)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for tag in (("w:top",), ("w:bottom",)) if double else (("w:bottom",),):
        el = OxmlElement(tag[0])
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    pPr.append(borders)
    return p


def _summary_line(doc, label: str, value: str, bold: bool = False):
    p = doc.add_paragraph()
    _tight(p)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.4), WD_TAB_ALIGNMENT.RIGHT)
    _run(p, label, 10, bold)
    _run(p, "\t" + value, 10, bold)
    return p


# ------------------------------------------------------------------ builder

COL_W = [Inches(2.7), Inches(0.95), Inches(0.95), Inches(0.75), Inches(0.95), Inches(0.75), Inches(0.85)]
HEADS = ["DESCRIPTION", "QUANTITY", "UNIT PRICE", "TAX", "RCV", "DEPREC.", "ACV"]


def build_estimate(cfg: EstimateConfig) -> dict:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ---- header (letterhead) ----
    header = doc.sections[0].header
    hp = header.paragraphs[0]
    if cfg.logo_path and os.path.exists(cfg.logo_path):
        run = hp.add_run()
        run.add_picture(cfg.logo_path, width=Inches(1.7))
        hp.add_run("  ")
    _run(hp, cfg.company_name, 11, bold=True)
    p2 = header.add_paragraph()
    _run(p2, f"{cfg.company_tagline}   |   {cfg.company_city}", 9)
    rule = header.add_paragraph()
    pPr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    el = OxmlElement("w:bottom")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:color"), "000000")
    borders.append(el)
    pPr.append(borders)

    # ---- footer ----
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.CENTER)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    _run(fp, cfg.estimate_name, 8)
    _run(fp, "\t" + cfg.footer_date, 8)
    r = fp.add_run("\tPage: ")
    r.font.name = FONT
    r.font.size = Pt(8)
    _add_page_number(fp)

    # ---- info block ----
    info = doc.add_table(rows=0, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_table_borders(info)

    def info_row(l1, v1, l2="", v2=""):
        row = info.add_row()
        for i, (txt, bold) in enumerate(((l1, True), (v1, False), (l2, True), (v2, False))):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            _tight(p)
            _run(p, txt, 10, bold)

    info_row("Property Owner:", cfg.owner, "Phone:", cfg.phone)
    prop = ", ".join(cfg.property_lines)
    info_row("Property:", prop)
    info_row("Estimator:", cfg.estimator_name, "Business:", cfg.company_city)
    info_row("Contractor Company:", cfg.contractor_company, "Business:", cfg.contractor_city)

    # ---- claim + dates ----
    p = doc.add_paragraph()
    _tight(p)
    p.paragraph_format.space_before = Pt(10)
    _run(p, "Claim Number: ", 10, True)
    _run(p, f"{cfg.claim_number}     ", 10)
    _run(p, "Policy Number: ", 10, True)
    _run(p, f"{cfg.policy_number}     ", 10)
    _run(p, "Type of Loss: ", 10, True)
    _run(p, cfg.type_of_loss, 10)

    p = doc.add_paragraph()
    _tight(p)
    _run(p, "Date of Loss: ", 10, True)
    _run(p, f"{cfg.date_of_loss}     ", 10)
    _run(p, "Date Inspected: ", 10, True)
    _run(p, f"{cfg.date_inspected}     ", 10)
    _run(p, "Date Entered: ", 10, True)
    _run(p, cfg.date_entered, 10)

    p = doc.add_paragraph()
    _tight(p)
    _run(p, "Price List: ", 10, True)
    _run(p, cfg.price_list, 10)
    p = doc.add_paragraph()
    _tight(p)
    _run(p, "Restoration/Service/Remodel", 10)
    p = doc.add_paragraph()
    _tight(p)
    _run(p, "Estimate: ", 10, True)
    _run(p, cfg.estimate_name, 10)

    # ---- structures ----
    built = []
    for idx, s in enumerate(cfg.structures):
        if idx > 0:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        _run(p, s["heading"], 11, True)

        # measurements block (two columns)
        meas = [mv for mv in s.get("measurements", []) if mv[0] or mv[1]]
        if meas:
            mt = doc.add_table(rows=0, cols=2)
            _no_table_borders(mt)
            for i in range(0, len(meas), 2):
                row = mt.add_row()
                for j in range(2):
                    if i + j < len(meas):
                        val, lab = meas[i + j]
                        cp = row.cells[j].paragraphs[0]
                        _tight(cp)
                        _run(cp, f"{val}  ", 9, True)
                        _run(cp, lab, 9)

        # items table
        t = doc.add_table(rows=1, cols=7)
        _no_table_borders(t)
        for i, hcell in enumerate(t.rows[0].cells):
            hcell.width = COL_W[i]
            hp2 = hcell.paragraphs[0]
            _tight(hp2)
            if i > 0:
                hp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(hp2, HEADS[i], 9, True)
            _cell_borders(hcell, bottom=True)

        base = tax = 0.0
        counter = 0
        for item in s["items"]:
            desc, qty, unit, price, matl_frac = item[:5]
            note = item[5] if len(item) > 5 else ""
            counter += 1
            line_base = qty * price
            line_tax = line_base * matl_frac * cfg.tax_rate
            rcv = line_base + line_tax
            base += line_base
            tax += line_tax
            row = t.add_row()
            vals = [
                f"{counter}. {desc}",
                f"{qty:,.2f} {unit}",
                money(price),
                money(line_tax),
                money(rcv),
                "0.00",
                money(rcv),
            ]
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.width = COL_W[i]
                cp = cell.paragraphs[0]
                _tight(cp)
                if i > 0:
                    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(cp, v, 9)
            if note:
                nrow = t.add_row()
                nc = nrow.cells[0]
                nc.merge(nrow.cells[6])
                cp = nc.paragraphs[0]
                _tight(cp)
                _run(cp, note, 8, italic=True)

        rcv_total = base + tax
        trow = t.add_row()
        tvals = [s.get("totals_label", "Totals:"), "", "", money(tax), money(rcv_total), "0.00", money(rcv_total)]
        for i, v in enumerate(tvals):
            cell = trow.cells[i]
            cell.width = COL_W[i]
            cp = cell.paragraphs[0]
            _tight(cp)
            if i > 0:
                cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(cp, v, 9, True)
            _cell_borders(cell, top=True)

        overhead = base * cfg.op_overhead_pct if cfg.include_op else 0.0
        profit = base * cfg.op_profit_pct if cfg.include_op else 0.0
        built.append({
            **s,
            "base": base,
            "tax": tax,
            "rcv": rcv_total,
            "overhead": overhead,
            "profit": profit,
            "summary_rcv": base + overhead + profit + tax,
        })

    grand_base = sum(s["base"] for s in built)
    grand_tax = sum(s["tax"] for s in built)
    grand_overhead = sum(s["overhead"] for s in built)
    grand_profit = sum(s["profit"] for s in built)
    grand_rcv = sum(s["summary_rcv"] for s in built)

    # ---- line item totals ----
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    _run(p, f"Line Item Totals: {cfg.estimate_name}", 10, True)
    _summary_line(doc, "Tax / RCV", f"{money(grand_tax)} / {money(sum(s['rcv'] for s in built))}")

    # ---- per-structure summaries ----
    for s in built:
        doc.add_page_break()
        p = doc.add_paragraph()
        _run(p, s.get("summary_heading", "Summary"), 11, True)
        _summary_line(doc, "Line Item Total", money(s["base"]))
        if cfg.include_op:
            _summary_line(doc, f"General Contractor Overhead ({round(cfg.op_overhead_pct*100)}%)", money(s["overhead"]))
            _summary_line(doc, f"General Contractor Profit ({round(cfg.op_profit_pct*100)}%)", money(s["profit"]))
        _summary_line(doc, f"Material Sales Tax ({cfg.tax_rate*100:.2f}%)", money(s["tax"]))
        _para_rule(doc, double=True)
        _summary_line(doc, "Replacement Cost Value", "$" + money(s["summary_rcv"]), bold=True)
        _summary_line(doc, "Net Claim", "$" + money(s["summary_rcv"]), bold=True)

    # ---- recap of taxes ----
    doc.add_page_break()
    p = doc.add_paragraph()
    _run(p, "Recap of Taxes", 11, True)
    rt = doc.add_table(rows=1, cols=3)
    _no_table_borders(rt)
    heads = ["", f"Material Sales Tax ({cfg.tax_rate*100:.2f}%)", f"Total Tax ({cfg.tax_rate*100:.2f}%)"]
    for i, hv in enumerate(heads):
        cp = rt.rows[0].cells[i].paragraphs[0]
        _tight(cp)
        if i > 0:
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(cp, hv, 9, True)
        _cell_borders(rt.rows[0].cells[i], bottom=True)
    for s in built:
        row = rt.add_row()
        for i, v in enumerate((s.get("recap_label", s["heading"]), money(s["tax"]), money(s["tax"]))):
            cp = row.cells[i].paragraphs[0]
            _tight(cp)
            if i > 0:
                cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _run(cp, v, 9)
    row = rt.add_row()
    for i, v in enumerate(("Total", money(grand_tax), money(grand_tax))):
        cp = row.cells[i].paragraphs[0]
        _tight(cp)
        if i > 0:
            cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(cp, v, 9, True)
        _cell_borders(row.cells[i], top=True)

    # ---- grand total ----
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    _run(p, "GRAND TOTAL, ALL STRUCTURES", 12, True)
    _summary_line(doc, "Line Item Total", money(grand_base), bold=True)
    if cfg.include_op:
        _summary_line(doc, f"General Contractor Overhead ({round(cfg.op_overhead_pct*100)}%)", money(grand_overhead))
        _summary_line(doc, f"General Contractor Profit ({round(cfg.op_profit_pct*100)}%)", money(grand_profit))
    _summary_line(doc, f"Material Sales Tax ({cfg.tax_rate*100:.2f}%)", money(grand_tax))
    _para_rule(doc, double=True)
    _summary_line(doc, "Replacement Cost Value (Net Claim)", "$" + money(grand_rcv), bold=True)

    os.makedirs(os.path.dirname(os.path.abspath(cfg.out_file)), exist_ok=True)
    doc.save(cfg.out_file)
    return {
        "out_file": cfg.out_file,
        "line_item_total": round(grand_base, 2),
        "material_sales_tax": round(grand_tax, 2),
        "overhead": round(grand_overhead, 2),
        "profit": round(grand_profit, 2),
        "replacement_cost_value": round(grand_rcv, 2),
    }


def _add_page_number(paragraph):
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)
